import sqlite3

import numpy as np
import pandas as pd
from datetime import date, datetime, timedelta
import os
from docx import Document
from docx.shared import Inches
import glob
import matplotlib.pyplot as plt


# %matplotlib inline


class ConnectSqlTable:
    def __init__(self, db):
        self.db = db

    def get_row_column(self, QUERY):
        # Connect to the database
        conn = sqlite3.connect(self.db)
        # Create a cursor
        cursor = conn.cursor()
        # Execute a query
        cursor.execute(QUERY)
        # Fetch the results
        results = cursor.fetchall()
        num_fields = len(cursor.description)
        field_names = [i[0] for i in cursor.description]
        # print(field_names, results)
        cursor.close()
        conn.close()
        return field_names, results


class PlotFigure:
    def __init__(self,
                 week,
                 date_format,
                 pic_format):
        self.week = week
        self.my_colors = list('rgbkymc')
        self.date_format = date_format
        self.pic_format = pic_format

    def figures(self, path=None, dt_coll=None, folder_name=None, rows_name=None, columns_name=None):
        df = pd.DataFrame(rows_name, columns=columns_name)
        # df = pd.read_csv(path)
        # df = df.iloc[::-1]
        df[dt_coll] = pd.to_datetime(df[dt_coll])
        # df = df.set_index(dt_coll)
        r = pd.date_range(start=df[dt_coll].min(), end=df[dt_coll].max())
        # print (r)
        df = df.set_index(dt_coll).reindex(r).fillna(np.nan).rename_axis(dt_coll)
        # print(df)
        df = df.iloc[::-1]
        # if the demo_folder_name directory is not present
        if not os.path.exists(folder_name):
            # then create it.
            os.makedirs(folder_name)
        # if Not sunday
        if datetime.today().weekday() != 6:
            # for 21 days only
            df1 = df[:df.first_valid_index() - pd.offsets.Day(self.week)]
            # saving file name
            save = df1.first_valid_index().strftime(self.date_format)
            # categorical Column
            cat_cols = df1.columns[df1.dtypes == "object"].tolist()
            for cat in cat_cols:
                # df1 = df[(df.T != 0).any()]
                # print(df1)
                # df1[cat].value_counts().plot(subplots=True, figsize=(10, 8), layout=(2, 2)).get_figure().savefig(
                #    folder_name + '/' + save + '_' + cat + '_' + str(self.week) + self.pic_format)
                df1[cat].value_counts().plot(kind='bar', y=cat, color=self.my_colors,
                                             use_index=True).get_figure().savefig(
                    folder_name + '/' + save + '_' + cat + '_' + str(self.week) + self.pic_format)
                # print(cat)
            # Numerical Column
            int_cols = df1.columns[df1.dtypes == "int64"].tolist()
            for value in int_cols:
                # df1.plot(subplots=True, figsize=(10, 8), layout=(2, 2)).get_figure().savefig(
                #    folder_name + '/' + save + '_' + value + '_' + str(self.week) + self.pic_format)
                df1.plot(kind='bar', y=value,
                         use_index=True).get_figure().savefig(
                    folder_name + '/' + save + '_' + value + '_' + str(self.week) + self.pic_format)
                # print(cat)
        else:  # elif datetime.today().weekday() == 6:
            try:
                weeks = 1
                while weeks <= self.week:
                    split_size = 7
                    # group the dataframe into 7 days each
                    dfs = [df[i:i + split_size:] for i in range(0, len(df), split_size)]
                    try:
                        for _, frame in enumerate(dfs):  # for each data frame in the list
                            # categorical Column
                            cat_cols = frame.columns[frame.dtypes == "object"].tolist()
                            # print(frame.first_valid_index())
                            save = frame.first_valid_index().strftime(self.date_format)
                            for cat in cat_cols:
                                frame[cat].value_counts().plot(kind='bar', y=cat, color=self.my_colors,
                                                               use_index=True).get_figure().savefig(
                                    folder_name + '/' + save + '_' + cat + 'week' + str(_) + self.pic_format)
                            # Numerical Column
                            int_cols = frame.columns[frame.dtypes == "int64"].tolist()
                            for value in int_cols:
                                frame.plot(kind='bar', y=value, use_index=True).get_figure().savefig(
                                    folder_name + '/' + save + '_' + value + 'week' + str(_) + self.pic_format)
                    except (TypeError, AttributeError):
                        pass
                    weeks += 1
                    print(weeks)
            except Exception as err:
                print(err)


class WordFile:
    def __init__(self, media_path, figure_path, date_format, pic_format):
        self.media: str = media_path
        self.figure: str = figure_path
        self.date_format: str = date_format
        self.pic_format = pic_format
        self.doc = Document()

    def add_content(
            self, date_check=None,
            para=None, slept=None, hours_slept=None,
            cigarette=None, sports=None, weather=None,
            grade=None, prognosis=None, doc_date_format=None):
        # print(date_check.strftime(self.date_format))
        mylist = [f for f in glob.glob(self.figure + '/' + date_check.strftime(self.date_format) + '*')]
        # print(mylist)
        today = datetime.today()
        yesterday = today - timedelta(days=1)

        # print(dats < yesterday)

        if date_check <= yesterday and len(mylist) > 0:
            self.doc.add_heading(date_check.strftime(doc_date_format), 1)
            if para:
                self.doc.add_paragraph(para)
            else:
                self.doc.add_paragraph('...')

            self.doc.add_paragraph('')
            self.doc.add_paragraph(f'Slept:{slept}')
            self.doc.add_paragraph(f'Slept time:{hours_slept}')
            self.doc.add_paragraph(f'Cigarette:{cigarette}')
            self.doc.add_paragraph(f'Sports:{sports}')
            self.doc.add_paragraph(f'Weather:{weather}')
            self.doc.add_paragraph(f'My Grade:{grade}')
            self.doc.add_paragraph(f'Prognosis:{prognosis}')

            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            # print(os.listdir(self.media))
            if os.path.exists(self.media):
                for filename in os.listdir(self.media):
                    if filename.startswith(date_check.strftime(self.date_format)) and filename.endswith(self.pic_format):
                        run.add_picture(os.path.join(self.media, filename), width=Inches(1.5), height=Inches(1.5))
            else:
                print('Media folder is not present in - Please create a folder :', self.media)

            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            if os.path.exists(self.figure):
                for filename in os.listdir(self.figure):
                    if filename.startswith(date_check.strftime(self.date_format)) and filename.endswith(self.pic_format):
                        run.add_picture(os.path.join(self.figure, filename), width=Inches(1.5), height=Inches(1.5))
            else:
                print('Media folder is not present in - Please create a folder :', self.figure)

            return self.doc

    def saveFile(
            self, rows_name=None, columns_name=None,
            row_predict=None,  columns_predict=None,
            doc_name=None, dt_coll=None, data=None):
        df = pd.DataFrame(rows_name, columns=columns_name)
        df[dt_coll] = pd.to_datetime(df[dt_coll])
        df = df.iloc[::-1]
        df = df.set_index(dt_coll)
        # print(df)

        df_temp = pd.DataFrame(row_predict, columns=columns_predict)
        df_temp[dt_coll] = pd.to_datetime(df_temp[dt_coll])
        df_temp = df_temp.iloc[::-1]
        df_temp = df_temp.set_index(dt_coll)
        # print(df_temp)

        # for index, rows in df.iterrows():
        for index, rows in df_temp.iterrows():
            # print(df.loc[index])
            # print(df.loc[index][data.sleep])
            # print('sleep time', df.loc[index][data.sleep])
            # print('Wake time', df.loc[index][data.wake_time])
            # print(df.loc[index][data.sleep].strptime("%H:%M")-df.loc[index][data.wake_time].strptime("%H:%M"))
            docs = self.add_content(
                    doc_date_format=data.doc_date_format,
                    date_check=df.loc[[index]].index.item(),  # [data.date],
                    # para=df.loc[index][data.para],  # place holder for the row para
                    grade=df.loc[index][data.answer],
                    prognosis=df_temp.loc[index][data.prediction],
                    slept=df.loc[index][data.sleep],
                    hours_slept=df.loc[index][data.wake_time],
                    cigarette=df.loc[index][data.smoke],
                    sports=df.loc[index][data.sport],
                    weather=df.loc[index][data.temp],)
            # print(docs)
            if docs is not None:
                docs.save(doc_name)
        print('FILE SUCCESSFULLY SAVED !')
