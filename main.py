# Importing all the required libraries



import mysql.connector as sql
import pandas as pd
from datetime import date,datetime,timedelta
import os
from docx import Document
from docx.shared import Inches
import glob

def connect_sql_table():
    # Connect to the database
    conn = sql.connect(
        host="localhost",
        user="root",
        password="Qwerty@#54321",
        database="predictions2.0"
    )

    # Create a cursor
    cursor = conn.cursor()

    # Execute a query
    cursor.execute("SELECT * FROM mytable")

    # Fetch the results
    results = cursor.fetchall()

    num_fields = len(cursor.description)
    field_names = [i[0] for i in cursor.description]

    cursor.close()
    conn.close()
    return field_names, results




def plot_figure(path=None,week=21,fold=None,rows_name=None,columns_name=None):
  df = pd.DataFrame(rows_name, columns=columns_name)
  #df = pd.read_csv(path)
  #print(df)
  df = df.iloc[::-1]
  #print(df)
  df['date'] = pd.to_datetime(df['date'])
  df = df.set_index('date')
  #print(df)
  my_colors = list('rgbkymc')
  if not os.path.exists(fold):
    # if the demo_folder directory is not present
    # then create it.
    os.makedirs(fold)
  if datetime.today().weekday() != 6:      #
    # for 21 days only
    #print(df.first_valid_index())

    df1 = df[:df.first_valid_index() - pd.offsets.Day(week)]
    #print(df1)
    save = df1.first_valid_index().strftime("%Y-%m-%d")
    # catorogical Column
    cat_cols = df1.columns[df1.dtypes == "object"].tolist()
    for cat in cat_cols:
      df1[cat].value_counts().plot(kind='bar', y=cat, color=my_colors, use_index = True).get_figure().savefig(fold+'/'+save+'_'+cat+'_'+str(week)+'.png')
      # print(cat)
    # Numerical Column
    int_cols = df1.columns[df1.dtypes == "int64"].tolist()
    for integ in int_cols:
      df1.plot(kind='bar', y=integ, use_index = True).get_figure().savefig(fold+'/'+save+'_'+integ+'_'+str(week)+'.png')
      # print(cat)
  else:   # elif datetime.today().weekday() == 6:
    try:
      print('here')
      weeks = 1
      while weeks <= check:
        split_size = 7
        # group the dataframe into 7 days each
        dfs = [df[i:i+split_size:] for i in range(0, len(df),split_size)]
        for _, frame in enumerate(dfs):   #for each data frame in the list
          # Catorogical Column
          cat_cols = frame.columns[frame.dtypes == "object"].tolist()
          save = frame.first_valid_index().strftime("%Y-%m-%d")
          for cat in cat_cols:
            frame[cat].value_counts().plot(kind='bar', y=cat, color=my_colors, use_index = True).get_figure().savefig(fold+'/'+save+'_'+cat+'week'+str(_)+'.png')
          # Numerical Column
          int_cols = frame.columns[frame.dtypes == "int64"].tolist()
          for integ in int_cols:
            frame.plot(kind='bar', y=integ, use_index = True).get_figure().savefig(fold+'/'+save+'_'+integ+'week'+str(_)+'.png')
          # print(cat)
          # if _ == 0:
          #  break
        weeks += 1
    except Exception as err:
      print(err.message, err.args)


def add_content(
        doc=None,
        dats=None,
        media=None,
        figure=None,
        para=None,
        slept=None,
        hours_slept=None,
        cigrattes=None,
        sports=None,
        weather=None,
        grade=None,
        prognosis=None):
    check = dats.strftime('%Y-%m-%d')
    mylist = [f for f in glob.glob(media + '/' + check + '*')]
    today = datetime.today()
    yesterday = today - timedelta(days=1)
    # print(dats < yesterday)
    if dats <= yesterday and len(mylist) > 0:
        ##########################
        doc.add_heading(dats.strftime("%A, %d. %b %Y"), 1)
        # doc.add_heading('Heading level 1', 1)
        if not para:
            doc.add_paragraph(para)
        else:
            doc.paragraph.add_run('...')
        # paragraph = doc.add_paragraph()
        doc.add_paragraph('')
        ##########################
        doc.add_paragraph(f'Slept:{slept}')
        ##########################
        doc.add_paragraph(f'Slept time:{hours_slept}')
        ##########################
        doc.add_paragraph(f'Cigrattes:{cigrattes}')
        ##########################
        doc.add_paragraph(f'Sports:{sports}')
        ##########################
        doc.add_paragraph(f'Weather:{weather}')
        ##########################
        doc.add_paragraph(f'My Grade:{grade}')
        ##########################
        doc.add_paragraph(f'Prognosis:{prognosis}')
        ##########################
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        for filename in os.listdir(media):
            if filename.startswith(check) and filename.endswith(".png"):
                run.add_picture(os.path.join(media, filename), width=Inches(1.5), height=Inches(1.5))
        ##########################
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        for filename in os.listdir(figure):
            if filename.startswith(check) and filename.endswith(".png"):
                run.add_picture(os.path.join(figure, filename), width=Inches(1.5), height=Inches(1.5))
        return doc


def save_doc(
        media=None,
        figure=None,
        rows_name=None,
        columns_name=None,
        path=None):
    today = datetime.today()
    yesterday = today - timedelta(days=1)

    df = pd.DataFrame(rows_name, columns=columns_name)
    df = df.iloc[::-1]
    #df = pd.read_csv(path)
    # df = df.set_index('date')
    df['date'] = pd.to_datetime(df['date'], format='%Y-%m-%d')
    doc = Document()
    for index, row in df.iterrows():
        # print(row[" user_bedtime"], row[" played_sport"])
        try:
            doc = add_content(
                doc=doc,
                dats=row['date'],
                media=media,
                figure=figure,
                # para=row['para'], # place holder for the row para
                slept=row[' hours_slept'],
                hours_slept=row[' hours_slept'],
                cigrattes=row[' smoked_cigeratte'],
                sports=row[' played_sport'],
                weather=row[' temp_yesterday'], )
            # grade=row['grade'], # graded place holder
            # prognosis=row['prognosis'] # graded place holder )
            doc.save('demomm.docx')
        except Exception as e:
            #  print(e)
            pass
        # print(doc)
    # print(doc)
    # doc.save('demo.docx')
    print('done')



# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    a,b = connect_sql_table()
    check = 21

    plot_figure(week=check, fold='figures', rows_name= b,columns_name= a)
    save_doc(
        media='/content/media',
        figure='/content/figures',
        rows_name= b,columns_name= a
    )











